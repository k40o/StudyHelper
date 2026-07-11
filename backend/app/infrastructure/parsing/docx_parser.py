"""Microsoft Word (.docx) parser, built on python-docx.

Word gives us paragraph *styles* ("Heading 1", "Title", "List Bullet"), which
are a reliable structural signal — far better than guessing from text.
"""
from __future__ import annotations

import re
from pathlib import Path

from ...domain.document import BlockType, ContentBlock, ParsedDocument, SourceType
from .base import DocumentParser, ParserError, derive_title

_HEADING_LEVEL_RE = re.compile(r"heading\s*(\d+)", re.IGNORECASE)


class DocxParser(DocumentParser):
    extensions = (".docx",)

    def parse(self, path: Path) -> ParsedDocument:
        try:
            from docx import Document  # imported lazily so the dep is optional
        except ImportError as exc:  # pragma: no cover
            raise ParserError("python-docx is not installed") from exc

        try:
            doc = Document(str(path))
        except Exception as exc:  # python-docx raises various errors
            raise ParserError(f"Could not open Word file {path}: {exc}") from exc

        blocks: list[ContentBlock] = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            block_type, level = self._classify_paragraph(para)
            blocks.append(ContentBlock(text=text, block_type=block_type, level=level))

        # Tables: flatten each row to "cell | cell | cell".
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                row_text = " | ".join(cells).strip(" |")
                if row_text:
                    blocks.append(ContentBlock(text=row_text, block_type=BlockType.TABLE))

        doc_title = (doc.core_properties.title or "").strip() or None
        first_heading = next((b.text for b in blocks if b.is_heading), None)
        return ParsedDocument(
            file_path=str(path),
            source_type=SourceType.DOCX,
            title=doc_title or derive_title(path, first_heading),
            blocks=blocks,
            metadata={
                "author": (doc.core_properties.author or "").strip(),
                "paragraphs": len(blocks),
            },
        )

    @staticmethod
    def _classify_paragraph(para) -> tuple[BlockType, int]:
        style_name = (para.style.name if para.style else "") or ""
        style_lower = style_name.lower()

        if style_lower == "title":
            return BlockType.HEADING, 1
        heading_match = _HEADING_LEVEL_RE.search(style_lower)
        if heading_match:
            return BlockType.HEADING, int(heading_match.group(1))
        if "list" in style_lower or para.text.strip().startswith(("•", "-", "*")):
            return BlockType.LIST_ITEM, 0
        return BlockType.PARAGRAPH, 0
